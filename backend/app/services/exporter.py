import os
import re
import html
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, KeepTogether, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfgen import canvas

# Canvas helper to handle page numbers ("Page X of Y") dynamically
class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_number(num_pages)
            canvas.Canvas.showPage(self)
        canvas.Canvas.save(self)

    def draw_page_number(self, page_count):
        self.saveState()
        self.setFont("Helvetica", 9)
        self.setFillColor(HexColor("#64748b")) # Slate 500
        
        # Header (on pages after the cover page)
        if self._pageNumber > 1:
            self.drawString(54, 750, "Autonomous AI Report Generation Agent")
            self.setStrokeColor(HexColor("#e2e8f0")) # Slate 200
            self.setLineWidth(0.5)
            self.line(54, 742, letter[0] - 54, 742)

        # Footer
        page_str = f"Page {self._pageNumber} of {page_count}"
        self.drawRightString(letter[0] - 54, 36, page_str)
        self.drawString(54, 36, "CONFIDENTIAL - SYSTEM GENERATED REPORT")
        self.restoreState()


def markdown_to_pdf(markdown_text: str, filename_or_buffer) -> None:
    """
    Renders markdown report to a PDF document with custom margins, colors, and layout.
    """
    doc = SimpleDocTemplate(
        filename_or_buffer,
        pagesize=letter,
        leftMargin=54,
        rightMargin=54,
        topMargin=72,
        bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    
    # Custom Palette Styling
    primary_color = "#0f172a"  # Slate 900
    secondary_color = "#2563eb"  # Blue 600
    text_color = "#334155"  # Slate 700
    
    # Custom Paragraph Styles
    title_style = ParagraphStyle(
        'ReportTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=24,
        leading=28,
        textColor=HexColor(primary_color),
        spaceAfter=15
    )
    
    h1_style = ParagraphStyle(
        'ReportH1',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=16,
        leading=20,
        textColor=HexColor(primary_color),
        spaceBefore=18,
        spaceAfter=8,
        keepWithNext=True
    )
    
    h2_style = ParagraphStyle(
        'ReportH2',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=13,
        leading=16,
        textColor=HexColor(secondary_color),
        spaceBefore=14,
        spaceAfter=6,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'ReportBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=HexColor(text_color),
        spaceAfter=8
    )

    list_style = ParagraphStyle(
        'ReportList',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=HexColor(text_color),
        leftIndent=15,
        spaceAfter=4
    )
    
    story = []
    
    # Parse Markdown lines simply
    lines = markdown_text.split('\n')
    for line in lines:
        line_strip = line.strip()
        if not line_strip:
            story.append(Spacer(1, 6))
            continue
            
        # Match headings
        if line_strip.startswith('# '):
            story.append(Paragraph(line_strip[2:], title_style))
            story.append(Spacer(1, 10))
        elif line_strip.startswith('## '):
            story.append(Paragraph(line_strip[3:], h1_style))
        elif line_strip.startswith('### '):
            story.append(Paragraph(line_strip[4:], h2_style))
        elif line_strip.startswith('* ') or line_strip.startswith('- '):
            clean_item = line_strip[2:]
            # Clean markdown bold/italic inside list items
            clean_item = clean_bold_markdown(clean_item)
            story.append(Paragraph(f"&bull; {clean_item}", list_style))
        else:
            clean_body = clean_bold_markdown(line_strip)
            story.append(Paragraph(clean_body, body_style))
            
    doc.build(story, canvasmaker=NumberedCanvas)


def clean_bold_markdown(text: str) -> str:
    """Escapes XML entities and replaces bold `**text**` and `*italic*` with HTML tags for ReportLab Paragraphs"""
    text = html.escape(text)
    # Replace **bold** with <b>bold</b>
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    # Replace *italic* with <i>italic</i>
    text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
    return text


def markdown_to_docx(markdown_text: str, filepath_or_buffer) -> None:
    """
    Converts markdown report to a Microsoft Word DOCX document.
    """
    doc = Document()
    
    # Base layout settings
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    primary_color = RGBColor(15, 23, 42)    # Slate 900
    secondary_color = RGBColor(37, 99, 235)  # Blue 600
    
    lines = markdown_text.split('\n')
    for line in lines:
        line_strip = line.strip()
        if not line_strip:
            continue
            
        if line_strip.startswith('# '):
            p = doc.add_heading(level=0)
            run = p.add_run(line_strip[2:])
            run.font.color.rgb = primary_color
            run.font.name = 'Arial'
            run.bold = True
        elif line_strip.startswith('## '):
            p = doc.add_heading(level=1)
            run = p.add_run(line_strip[3:])
            run.font.color.rgb = primary_color
            run.font.name = 'Arial'
        elif line_strip.startswith('### '):
            p = doc.add_heading(level=2)
            run = p.add_run(line_strip[4:])
            run.font.color.rgb = secondary_color
            run.font.name = 'Arial'
        elif line_strip.startswith('* ') or line_strip.startswith('- '):
            clean_item = line_strip[2:]
            p = doc.add_paragraph(style='List Bullet')
            parse_styled_docx_run(p, clean_item)
        else:
            p = doc.add_paragraph()
            parse_styled_docx_run(p, line_strip)
            
    doc.save(filepath_or_buffer)


def parse_styled_docx_run(paragraph, text: str) -> None:
    """Helper to parse bold/italic formatting in raw text and add to docx paragraph"""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Calibri'
        else:
            subparts = re.split(r'(\*.*?\*)', part)
            for subpart in subparts:
                if subpart.startswith('*') and subpart.endswith('*'):
                    run = paragraph.add_run(subpart[1:-1])
                    run.italic = True
                    run.font.name = 'Calibri'
                else:
                    run = paragraph.add_run(subpart)
                    run.font.name = 'Calibri'
                    run.font.color.rgb = RGBColor(51, 65, 85) # Slate 700
